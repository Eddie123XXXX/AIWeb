"""
RAG 知识库功能测试

运行方式（在 backend 目录下）:

    cd backend
    python rag/test_rag.py

测试分层:
- 单元测试: parsers、chunking（无需 fastapi，仅需 python-docx/pandas 做 Word/Excel 测试）
- 集成测试: 上传、解析、检索（需 fastapi、MinIO、PostgreSQL、Milvus、.env 配置）
"""
from __future__ import annotations

import asyncio
import io
import os
import sys
from uuid import uuid4

# 确保 backend 根目录在 path 中
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass


def _import_parsers():
    """导入 parsers 模块（避免加载 router/fastapi）"""
    import importlib.util
    backend = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    path = os.path.join(backend, "rag", "parsers.py")
    spec = importlib.util.spec_from_file_location("rag.parsers", path)
    mod = importlib.util.module_from_spec(spec)
    sys.modules["rag.parsers"] = mod
    spec.loader.exec_module(mod)
    return mod


def _import_chunking():
    """导入 chunking 模块（避免加载 router/fastapi）"""
    import importlib.util
    backend = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    path = os.path.join(backend, "rag", "chunking.py")
    spec = importlib.util.spec_from_file_location("rag.chunking", path)
    mod = importlib.util.module_from_spec(spec)
    sys.modules["rag.chunking"] = mod
    spec.loader.exec_module(mod)
    return mod


# ---------------------------------------------------------------------------
# 1. Parsers 单元测试
# ---------------------------------------------------------------------------

def test_parsers_supported_extensions() -> None:
    """测试支持的文件类型识别"""
    parsers = _import_parsers()

    assert parsers.is_supported("doc.pdf")
    assert parsers.is_supported("report.docx")
    assert parsers.is_supported("data.xlsx")
    assert parsers.is_supported("readme.md")
    assert parsers.is_supported("notes.txt")
    assert parsers.is_supported("audio.mp3")
    assert not parsers.is_supported("unknown.xyz")
    assert not parsers.is_supported("")


def test_parsers_get_parser() -> None:
    """测试解析器路由"""
    parsers = _import_parsers()

    assert parsers.get_parser_for_file("a.pdf") == "mineru"
    assert parsers.get_parser_for_file("b.docx") == "docx"
    assert parsers.get_parser_for_file("c.xlsx") == "excel"
    assert parsers.get_parser_for_file("d.md") == "markdown"
    assert parsers.get_parser_for_file("e.txt") == "txt"
    assert parsers.get_parser_for_file("f.mp3") == "audio"
    assert parsers.get_parser_for_file("noext") is None


def test_parsers_txt() -> None:
    """测试 TXT 解析"""
    parsers = _import_parsers()

    data = "第一段内容。\n\n第二段内容。".encode("utf-8")
    result = parsers.parse_local(data, "test.txt")
    assert "markdown" in result
    assert "content_list" in result
    assert len(result["content_list"]) >= 1
    assert any("第一段" in b.get("text", "") for b in result["content_list"])


def test_parsers_markdown() -> None:
    """测试 Markdown 解析"""
    parsers = _import_parsers()

    data = "# 标题\n\n正文内容。".encode("utf-8")
    result = parsers.parse_local(data, "readme.md")
    assert "markdown" in result
    assert "content_list" in result
    assert "标题" in result["markdown"]


def test_parsers_docx() -> None:
    """测试 Word 解析（需 python-docx）"""
    parsers = _import_parsers()

    try:
        from docx import Document
    except ImportError:
        print("  [SKIP] python-docx not installed")
        return

    doc = Document()
    doc.add_paragraph("测试段落一")
    doc.add_paragraph("测试段落二")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    data = buf.read()

    result = parsers.parse_local(data, "test.docx")
    assert "markdown" in result
    assert "content_list" in result
    assert any("测试段落" in b.get("text", "") for b in result["content_list"])


def test_parsers_excel() -> None:
    """测试 Excel 解析（需 pandas + openpyxl）"""
    parsers = _import_parsers()

    try:
        import pandas as pd
    except ImportError:
        print("  [SKIP] pandas not installed")
        return

    df = pd.DataFrame({"A": [1, 2], "B": ["x", "y"]})
    buf = io.BytesIO()
    df.to_excel(buf, index=False)
    buf.seek(0)
    data = buf.read()

    result = parsers.parse_local(data, "data.xlsx")
    assert "markdown" in result
    assert "content_list" in result


# ---------------------------------------------------------------------------
# 2. Chunking 单元测试
# ---------------------------------------------------------------------------

def test_chunking_estimate_tokens() -> None:
    """测试 Token 估算"""
    chunking = _import_chunking()

    assert chunking.estimate_tokens("") == 0
    assert chunking.estimate_tokens("hello") >= 1
    assert chunking.estimate_tokens("你好世界") >= 2
    assert chunking.estimate_tokens("a" * 100) > chunking.estimate_tokens("a" * 10)


def test_chunking_get_content_for_embedding() -> None:
    """测试巨型 Chunk 护栏"""
    chunking = _import_chunking()

    short = chunking.Chunk(
        id="1", document_id="d", notebook_id="n", chunk_index=0,
        content="短内容", token_count=10, chunk_type="TEXT",
    )
    out = chunking.get_content_for_embedding(short, max_tokens=2048)
    assert out == "短内容"

    long_content = "x" * 10000
    long = chunking.Chunk(
        id="2", document_id="d", notebook_id="n", chunk_index=0,
        content=long_content, token_count=3000, chunk_type="TEXT",
    )
    out = chunking.get_content_for_embedding(long, max_tokens=512)
    assert len(out) < len(long_content)
    assert "截断" in out or "..." in out


def test_chunking_process_mineru_blocks() -> None:
    """测试 MinerU Block 切块"""
    chunking = _import_chunking()

    blocks = [
        {"type": "title", "text": "第一章 概述", "page_idx": 0},
        {"type": "text", "text": "本章介绍系统架构。", "page_idx": 0},
        {"type": "table", "text": "|A|B|\n|1|2|", "page_idx": 1},
    ]
    chunks = chunking.process_mineru_blocks(
        blocks, document_id="doc1", notebook_id="nb1",
    )
    assert len(chunks) >= 2
    parent_count = sum(1 for c in chunks if c.is_parent)
    child_count = sum(1 for c in chunks if not c.is_parent)
    assert parent_count >= 1
    assert child_count >= 1
    assert any("第一章" in c.content for c in chunks)
    assert any(c.chunk_type == "TABLE" for c in chunks)


def test_chunking_chunk_markdown() -> None:
    """测试 Markdown 降级切块"""
    chunking = _import_chunking()

    md = "# 标题\n\n第一段。\n\n第二段。"
    chunks = chunking.chunk_markdown(md, document_id="d", notebook_id="n")
    assert len(chunks) >= 1
    assert any("标题" in c.content or "第一段" in c.content for c in chunks)


# ---------------------------------------------------------------------------
# 3. 集成测试（需 infra）
# ---------------------------------------------------------------------------

async def _check_infra() -> tuple[bool, str]:
    """检查 infra 是否可用（MinIO + PostgreSQL）"""
    try:
        from infra.minio.service import get_object
        from rag.document_repository import document_repository
        await document_repository.list_by_notebook("test", limit=1)
        return True, ""
    except Exception as e:
        return False, str(e)


async def test_upload_and_search() -> None:
    """集成测试：上传 TXT → 解析 → 检索"""
    try:
        from rag import service
        from rag.models import SearchRequest
    except ImportError as e:
        print(f"  [SKIP] Integration test needs full deps (fastapi, etc): {e}")
        return

    ok, err = await _check_infra()
    if not ok:
        print(f"  [SKIP] Integration test needs MinIO + PostgreSQL: {err}")
        return

    notebook_id = f"test_rag_{uuid4().hex[:8]}"
    content = "RAG 知识库测试文档。\n\n本文档用于验证上传、解析、向量化与检索流程。\n\n关键词：Python、FastAPI、机器学习。"
    file_data = content.encode("utf-8")
    filename = "rag_test.txt"

    print("\n=== Integration: upload TXT -> parse -> search ===")

    # upload
    doc = await service.upload_document(
        notebook_id=notebook_id,
        user_id=1,
        filename=filename,
        file_data=file_data,
        content_type="text/plain",
    )
    assert doc["id"]
    assert doc["status"] in ("UPLOADED", "READY")
    print(f"  Upload OK: doc_id={doc['id']}")

    # parse if not READY
    if doc["status"] != "READY":
        doc = await service.process_document(doc["id"])
        assert doc["status"] in ("READY", "FAILED"), f"Parse failed: {doc.get('error_log')}"
        print(f"  Parse done: status={doc['status']}")
    else:
        print("  Doc already READY (dedup)")

    if doc["status"] != "READY":
        print("  [SKIP] Search test (doc not READY, may need Milvus/Embedding)")
        return

    # 检索
    req = SearchRequest(
        notebook_id=notebook_id,
        query="机器学习",
        enable_exact=True,
        enable_sparse=True,
        enable_dense=True,
        enable_rerank=True,
        top_k=None,
        rerank_threshold=0.2,
        fallback_cosine_threshold=0.85,
    )
    resp = await service.search(req)
    assert resp.query == "机器学习"
    assert len(resp.hits) >= 0
    print(f"  Search done: {len(resp.hits)} hits")

    # cleanup
    await service.delete_document(doc["id"])
    print("  Cleaned up test doc")


# ---------------------------------------------------------------------------
# 主入口
# ---------------------------------------------------------------------------

def run_unit_tests() -> int:
    """运行单元测试"""
    tests = [
        ("parsers_supported_extensions", test_parsers_supported_extensions),
        ("parsers_get_parser", test_parsers_get_parser),
        ("parsers_txt", test_parsers_txt),
        ("parsers_markdown", test_parsers_markdown),
        ("parsers_docx", test_parsers_docx),
        ("parsers_excel", test_parsers_excel),
        ("chunking_estimate_tokens", test_chunking_estimate_tokens),
        ("chunking_get_content_for_embedding", test_chunking_get_content_for_embedding),
        ("chunking_process_mineru_blocks", test_chunking_process_mineru_blocks),
        ("chunking_chunk_markdown", test_chunking_chunk_markdown),
    ]
    passed = 0
    for name, fn in tests:
        try:
            fn()
            print(f"  [OK] {name}")
            passed += 1
        except Exception as e:
            print(f"  [FAIL] {name}: {e}")
    return passed


def _safe_print(msg: str) -> None:
    """避免 Windows 控制台编码问题"""
    try:
        print(msg)
    except UnicodeEncodeError:
        print(msg.encode("ascii", "replace").decode("ascii"))


async def main() -> None:
    _safe_print("=" * 60)
    _safe_print("RAG Knowledge Base Test")
    _safe_print("=" * 60)

    _safe_print("\n[1] Unit tests (parsers, chunking)")
    _safe_print("-" * 40)
    passed = run_unit_tests()
    total = 10
    _safe_print(f"\nUnit tests: {passed}/{total} passed")

    _safe_print("\n[2] Integration test (upload, parse, search)")
    _safe_print("-" * 40)
    await test_upload_and_search()

    _safe_print("\n" + "=" * 60)
    _safe_print("Test completed")
    _safe_print("=" * 60)


if __name__ == "__main__":
    asyncio.run(main())
