"""
RAG 多格式文档解析器

支持: PDF (MinerU), Word, Excel, Markdown, TXT, MP3 等常用格式。
统一输出格式: {"markdown": str, "content_list": [{"type": str, "text": str, "page_idx": int}, ...]}
"""
from __future__ import annotations

import io
import logging
import os
from typing import Any

logger = logging.getLogger("rag.parsers")

# 支持的文件类型 (扩展名 -> 解析器名)
SUPPORTED_EXTENSIONS = {
    ".pdf": "mineru",
    ".docx": "docx",
    ".xlsx": "excel",
    ".xls": "excel",
    ".md": "markdown",
    ".markdown": "markdown",
    ".txt": "txt",
    ".text": "txt",
    ".mp3": "audio",
    ".wav": "audio",
    ".m4a": "audio",
    ".webm": "audio",
}


def _get_extension(filename: str) -> str:
    """从文件名提取扩展名 (小写)"""
    if not filename:
        return ""
    return os.path.splitext(filename)[1].lower()


def get_parser_for_file(filename: str) -> str | None:
    """根据文件名返回解析器名，不支持则返回 None"""
    ext = _get_extension(filename)
    return SUPPORTED_EXTENSIONS.get(ext)


def is_supported(filename: str) -> bool:
    """检查文件类型是否支持"""
    return get_parser_for_file(filename) is not None


# ---------------------------------------------------------------------------
# 各格式解析器 (同步，返回统一格式)
# ---------------------------------------------------------------------------

def _parse_txt(data: bytes) -> dict[str, Any]:
    """纯文本解析"""
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = data.decode("gbk")
        except UnicodeDecodeError:
            text = data.decode("utf-8", errors="replace")
    if not text.strip():
        return {"markdown": "", "content_list": []}
    blocks = [{"type": "text", "text": p.strip(), "page_idx": 0} for p in text.split("\n\n") if p.strip()]
    return {"markdown": text, "content_list": blocks if blocks else [{"type": "text", "text": text, "page_idx": 0}]}


def _parse_markdown(data: bytes) -> dict[str, Any]:
    """Markdown 解析 (与 txt 类似，保留结构)"""
    result = _parse_txt(data)
    # Markdown 可直接作为 markdown 字段，content_list 按段落
    return result


def _parse_pdf_local(data: bytes) -> dict[str, Any]:
    """PDF 本地解析 (pdfplumber)，MinerU 不可用时的降级方案"""
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("PDF 本地解析需安装 pdfplumber: pip install pdfplumber")

    blocks: list[dict[str, Any]] = []
    parts: list[str] = []

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page_idx, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text and text.strip():
                for para in text.strip().split("\n\n"):
                    p = para.strip()
                    if p:
                        blocks.append({"type": "text", "text": p, "page_idx": page_idx})
                        parts.append(p)
            # 表格
            tables = page.extract_tables()
            for table in tables or []:
                if table:
                    rows = ["| " + " | ".join(str(c or "") for c in row) + " |" for row in table]
                    table_md = "\n".join(rows)
                    if table_md.strip():
                        blocks.append({"type": "table", "text": table_md, "page_idx": page_idx})
                        parts.append(table_md)

    markdown = "\n\n".join(parts) if parts else ""
    return {"markdown": markdown, "content_list": blocks}


def _parse_docx(data: bytes) -> dict[str, Any]:
    """Word 文档解析 (python-docx)"""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("解析 Word 需安装 python-docx: pip install python-docx")

    doc = Document(io.BytesIO(data))
    blocks: list[dict[str, Any]] = []
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if "heading" in style or "title" in style:
            block_type = "title"
        else:
            block_type = "text"
        blocks.append({"type": block_type, "text": text, "page_idx": 0})
        parts.append(text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        table_md = "\n".join(rows)
        if table_md.strip():
            blocks.append({"type": "table", "text": table_md, "page_idx": 0})
            parts.append(table_md)

    markdown = "\n\n".join(parts) if parts else ""
    return {"markdown": markdown, "content_list": blocks}


def _parse_excel(data: bytes) -> dict[str, Any]:
    """Excel 解析 (pandas)"""
    try:
        import pandas as pd
    except ImportError:
        raise RuntimeError("解析 Excel 需安装 pandas: pip install pandas")

    blocks: list[dict[str, Any]] = []
    parts: list[str] = []

    try:
        xl = pd.ExcelFile(io.BytesIO(data))
        for sheet_name in xl.sheet_names:
            df = pd.read_excel(xl, sheet_name=sheet_name)
            if df.empty:
                continue
            blocks.append({
                "type": "title",
                "text": f"表: {sheet_name}",
                "page_idx": 0,
            })
            parts.append(f"## {sheet_name}")

            try:
                table_md = df.to_markdown(index=False)
            except Exception:
                table_md = df.to_string(index=False)
            blocks.append({"type": "table", "text": table_md, "page_idx": 0})
            parts.append(table_md)
    except Exception as e:
        logger.warning(f"[RAG] Excel 解析异常: {e}")
        raise RuntimeError(f"Excel 解析失败: {e}") from e

    markdown = "\n\n".join(parts) if parts else ""
    return {"markdown": markdown, "content_list": blocks}


def _parse_audio(data: bytes, filename: str) -> dict[str, Any]:
    """音频转文字 (OpenAI Whisper API)"""
    try:
        from openai import OpenAI
    except ImportError:
        raise RuntimeError("解析音频需安装 openai: pip install openai")

    api_key = os.getenv("OPENAI_API_KEY") or os.getenv("QWEN_API_KEY", "")
    if not api_key.strip():
        raise RuntimeError("音频转写需配置 OPENAI_API_KEY 或 QWEN_API_KEY")

    base_url = os.getenv("OPENAI_API_BASE") or None
    if os.getenv("QWEN_API_KEY"):
        base_url = os.getenv("QWEN_API_BASE", "https://dashscope.aliyuncs.com/compatible-mode/v1")

    client = OpenAI(api_key=api_key, base_url=base_url) if base_url else OpenAI(api_key=api_key)

    ext = _get_extension(filename)
    # Whisper 支持 mp3, mp4, mpeg, mpga, m4a, wav, webm
    suffix = ext if ext else ".mp3"
    file_obj = io.BytesIO(data)
    file_obj.name = f"audio{suffix}"

    try:
        resp = client.audio.transcriptions.create(
            model="whisper-1",
            file=file_obj,
            response_format="text",
        )
        text = resp if isinstance(resp, str) else getattr(resp, "text", str(resp))
    except Exception as e:
        logger.warning(f"[RAG] Whisper 转写失败: {e}")
        raise RuntimeError(f"音频转写失败: {e}") from e

    if not text.strip():
        return {"markdown": "", "content_list": []}
    blocks = [{"type": "text", "text": p.strip(), "page_idx": 0} for p in text.split("\n\n") if p.strip()]
    if not blocks:
        blocks = [{"type": "text", "text": text, "page_idx": 0}]
    return {"markdown": text, "content_list": blocks}


def parse_local(data: bytes, filename: str) -> dict[str, Any]:
    """
    本地解析：根据扩展名选择解析器。

    PDF 可由 MinerU 远程解析，MinerU 不可用时由 parse_pdf_local 降级。
    """
    parser = get_parser_for_file(filename)
    if not parser:
        raise ValueError(f"不支持的文件类型: {filename}")

    if parser == "mineru":
        return _parse_pdf_local(data)
    if parser == "txt":
        return _parse_txt(data)
    if parser == "markdown":
        return _parse_markdown(data)
    if parser == "docx":
        return _parse_docx(data)
    if parser == "excel":
        return _parse_excel(data)
    if parser == "audio":
        return _parse_audio(data, filename)

    raise ValueError(f"未实现的解析器: {parser}")
